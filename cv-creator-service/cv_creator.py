# -*- coding: utf-8 -*-
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches, Cm
from docx.shared import Pt

def template1(json_object):
    # initiate a new word document with custom margins
    document = Document()
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

    obj_styles = document.styles
    obj_charstyle = obj_styles.add_style('CommentsStyle', WD_STYLE_TYPE.CHARACTER)
    obj_font = obj_charstyle.font
    obj_font.size = Pt(12)
    obj_font.name = 'Times New Roman'

    # add user info to document on upper right-hand corner
    name = json_object['fullname']
    address = json_object['address']
    telephone = json_object['telephone']
    email = json_object['email']

    paragraph = document.add_paragraph()
    paragraph.add_run('Location: ').bold = True
    paragraph.add_run(address)
    run = paragraph.add_run()
    run.add_break()
    paragraph.add_run("Telephone: ").bold = True
    paragraph.add_run(telephone) # replace with input
    run = paragraph.add_run()
    run.add_break()
    paragraph.add_run("Email: ").bold = True
    paragraph.add_run(email)
    street_address_p_format = paragraph.paragraph_format
    street_address_p_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    certificates_achievements = ""

    # add name to document
    document.add_heading(name, 0)

    # profile section
    document.add_heading('Professional Profile', level=1)
    professional_profile_p = document.add_paragraph(json_object['strengths'])

    # education section
    document.add_heading('Education', level=1)
    education_items = int(len(json_object['education']))

    # loop through each item for education
    for i in range(education_items):
        school = json_object['education'][i]['school']
        period = json_object['education'][i]['period']
        description = json_object['education'][i]['description']
        education_p = document.add_paragraph(style='List Bullet')
        education_p.add_run(school + " (" + period + ")").bold = True
        run = education_p.add_run()
        run.add_break()
        education_p.add_run(description)

    # work experience section
    document.add_heading('Work Experience', level=1)
    work_experience_items = int(len(json_object['workexp']))
    
    # loop through each item for work experience 
    for i in range(work_experience_items):
        company = json_object['workexp'][i]['company']
        period = json_object['workexp'][i]['period']
        description = json_object['workexp'][i]['description']
        work_experience_p = document.add_paragraph(style='List Bullet')
        work_experience_p.add_run(company + " (" + period + ")").bold = True
        run = work_experience_p.add_run()
        run.add_break()
        work_experience_p.add_run(description)

        # volunteering section
    document.add_heading('Volunteering', level=1)
    volunteering_items = int(len(json_object['volunteering']))

    # loop through each item for volunteering experience 
    for i in range(volunteering_items):
        organisation = json_object['volunteering'][i]['org']
        period = json_object['volunteering'][i]['period']
        description = json_object['volunteering'][i]['description']
        volunteering_p = document.add_paragraph(style='List Bullet')
        volunteering_p.add_run(organisation + " (" + period + ")").bold = True
        run = volunteering_p.add_run()
        run.add_break()
        volunteering_p.add_run(description)
        
    # add skills section
    document.add_heading('Skills', level=1)
    skills_p = document.add_paragraph(style='List Bullet')
    skills_p.add_run(json_object['skills'])

    # add achievments/accolades/certificates section
    document.add_heading('Achievments/Accolades/Certificates', level=1)
    certificates_achievements_p = document.add_paragraph(style='List Bullet')
    certificates_achievements_p.add_run(json_object['profqual']['qual'] + " (" + json_object['profqual']['date'] + ")").bold = True

    # document.add_picture('butterfly.png', width=Inches(3))

    # # document.add_page_break()

    # document.save('demo.docx')

    return document

# if __name__ == "__main__":
#     template1()